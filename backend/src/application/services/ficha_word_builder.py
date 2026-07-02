"""
Generador de fichas Word institucionales con python-docx.
Replica el diseño visual del PDF (colores, portada, tablas y secciones).
"""

import os
from io import BytesIO
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from src.application.services.ficha_pdf_builder import SECTION_ICONS

RGB_PRIMARY = RGBColor(0x0F, 0x4A, 0xA0)
RGB_ACCENT = RGBColor(0x1D, 0x9E, 0x75)
RGB_TEXT = RGBColor(0x1E, 0x29, 0x3B)
RGB_MUTED = RGBColor(0x64, 0x74, 0x8B)
HEX_LIGHT = 'F0F6FF'
HEX_BORDER = 'E2E8F0'
HEX_WHITE = 'FFFFFF'


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, value in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), value.get('val', 'single'))
        element.set(qn('w:sz'), value.get('sz', '4'))
        element.set(qn('w:color'), value.get('color', HEX_BORDER))
        borders.append(element)
    tc_pr.append(borders)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_horizontal_rule(paragraph, color_hex: str = '0F4AA0') -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


class FichaWordBuilder:
    def __init__(self, ficha: Dict[str, Any]):
        self.ficha = ficha

    def build(self) -> bytes:
        doc = Document()
        self._configure_page(doc)
        self._configure_header_footer(doc)
        self._cover_block(doc)
        doc.add_page_break()

        for section in self.ficha.get('sections') or []:
            self._section_block(doc, section)

        doc.add_page_break()
        self._images_block(doc)

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _configure_page(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(10)
        normal.font.color.rgb = RGB_TEXT

    def _configure_header_footer(self, doc: Document) -> None:
        section = doc.sections[0]
        sistema = self.ficha.get('sistema', 'Pelileo Turismo')
        tipo_label = self.ficha.get('tipo_label', '')
        generado = self.ficha.get('generado_en', '')

        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        if header_para.text:
            header_para.clear()
        left = header_para.add_run(sistema)
        left.bold = True
        left.font.size = Pt(9)
        left.font.color.rgb = RGB_PRIMARY
        header_para.add_run('\t')
        right = header_para.add_run(tipo_label)
        right.font.size = Pt(8)
        right.font.color.rgb = RGB_MUTED
        header_para.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)

        rule = header.add_paragraph()
        _add_horizontal_rule(rule, '0F4AA0')

        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if footer_para.text:
            footer_para.clear()
        gen_run = footer_para.add_run(f'Generado: {generado}')
        gen_run.font.size = Pt(8)
        gen_run.font.color.rgb = RGB_MUTED
        footer_para.add_run('\t')
        page_label = footer_para.add_run('Página ')
        page_label.font.size = Pt(8)
        page_label.font.color.rgb = RGB_MUTED
        _add_page_number(footer_para)
        footer_para.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)

        accent_rule = footer.add_paragraph()
        _add_horizontal_rule(accent_rule, '1D9E75')

    def _paragraph(
        self,
        doc: Document,
        text: str,
        *,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size: int = 10,
        color: RGBColor = RGB_TEXT,
        bold: bool = False,
        space_after: int = 6,
        space_before: int = 0,
    ):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        return p

    def _cover_block(self, doc: Document) -> None:
        logo_path = self.ficha.get('logo_path')
        if logo_path and os.path.isfile(logo_path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(logo_path, width=Inches(1.4))
                doc.add_paragraph('')
            except Exception:
                pass

        self._paragraph(
            doc,
            self.ficha.get('sistema', 'Pelileo Turismo'),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=22,
            color=RGB_PRIMARY,
            bold=True,
            space_after=8,
        )
        self._paragraph(
            doc,
            self.ficha.get('eslogan', ''),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=12,
            color=RGB_MUTED,
            space_after=4,
        )
        doc.add_paragraph('')
        self._paragraph(doc, 'FICHA INSTITUCIONAL', align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=RGB_MUTED)
        self._paragraph(
            doc,
            self.ficha.get('tipo_label', ''),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=12,
            color=RGB_MUTED,
            space_after=12,
        )
        self._paragraph(
            doc,
            self.ficha.get('titulo', ''),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=18,
            color=RGB_TEXT,
            bold=True,
            space_after=8,
        )
        self._paragraph(
            doc,
            f"Documento generado el {self.ficha.get('generado_en', '')}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=12,
            color=RGB_MUTED,
        )

    def _section_heading(self, doc: Document, title: str) -> None:
        icon = SECTION_ICONS.get(title, '▪')
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f'{icon}  {title}')
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGB_PRIMARY

    def _fields_table(self, doc: Document, fields: List[Tuple[str, str]]) -> None:
        if not fields:
            return

        table = doc.add_table(rows=len(fields), cols=2)
        table.autofit = False
        table.allow_autofit = False

        for idx, (label, value) in enumerate(fields):
            row = table.rows[idx]
            row.height = Cm(0.65)
            bg = HEX_LIGHT if idx % 2 else HEX_WHITE

            label_cell = row.cells[0]
            value_cell = row.cells[1]
            label_cell.width = Cm(5.5)
            value_cell.width = Cm(11.5)

            for cell in (label_cell, value_cell):
                _set_cell_shading(cell, bg)
                _set_cell_border(
                    cell,
                    bottom={'val': 'single', 'sz': '4', 'color': HEX_BORDER},
                )

            label_para = label_cell.paragraphs[0]
            label_para.paragraph_format.space_after = Pt(2)
            label_para.paragraph_format.space_before = Pt(2)
            label_run = label_para.add_run(str(label))
            label_run.bold = True
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = RGB_MUTED

            value_para = value_cell.paragraphs[0]
            value_para.paragraph_format.space_after = Pt(2)
            value_para.paragraph_format.space_before = Pt(2)
            value_text = str(value).replace('\n', '\n')
            value_run = value_para.add_run(value_text)
            value_run.font.size = Pt(10)
            value_run.font.color.rgb = RGB_TEXT

        doc.add_paragraph('')

    def _text_block(self, doc: Document, label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(str(label))
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGB_MUTED

        body = doc.add_paragraph(str(value))
        body.paragraph_format.space_after = Pt(8)
        for run in body.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGB_TEXT

    def _section_block(self, doc: Document, section: Dict[str, Any]) -> None:
        self._section_heading(doc, section['title'])
        if section.get('fields'):
            self._fields_table(doc, section['fields'])
        for label, value in section.get('text_blocks') or []:
            self._text_block(doc, label, value)

    def _images_block(self, doc: Document) -> None:
        images = self.ficha.get('images') or []
        self._section_heading(doc, 'Galería de imágenes')

        if not images:
            self._paragraph(doc, 'No registrado', size=10, color=RGB_TEXT)
            return

        for img_meta in images:
            path = img_meta.get('path')
            caption = img_meta.get('titulo') or 'Imagen'
            if img_meta.get('descripcion'):
                caption = f"{caption} — {img_meta['descripcion']}"

            if path and os.path.isfile(path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(8)
                    run = p.add_run()
                    run.add_picture(path, width=Inches(5.5))
                    cap = doc.add_paragraph(caption)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_after = Pt(12)
                    for run in cap.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGB_MUTED
                except Exception:
                    self._paragraph(
                        doc,
                        f"{caption}: no se pudo incluir el archivo.",
                        size=10,
                        color=RGB_MUTED,
                    )
            else:
                self._paragraph(doc, f'• {caption}', size=10, color=RGB_MUTED)
